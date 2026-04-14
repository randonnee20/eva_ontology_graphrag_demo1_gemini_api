"""
pipeline/converter.py
원본 문서를 텍스트로 변환: pdf, docx, hwp, hwpx
"""

import os
import re
import subprocess
import zipfile
import logging
from pathlib import Path

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

_SOFFICE_CANDIDATES = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "soffice",
]


def _find_soffice() -> str:
    for candidate in _SOFFICE_CANDIDATES:
        if candidate == "soffice":
            return candidate
        if Path(candidate).exists():
            return candidate
    return "soffice"


def convert_to_text(input_path: str, output_path: str) -> str:
    ext = Path(input_path).suffix.lower().lstrip(".")
    converters = {"pdf": _from_pdf, "docx": _from_docx, "hwpx": _from_hwpx, "hwp": _from_hwp}

    if ext not in converters:
        raise ValueError(f"지원하지 않는 형식: .{ext}")

    text = converters[ext](input_path)

    if not text or len(text.strip()) < 10:
        logger.warning(f"텍스트 추출 결과 짧음: {input_path}")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(text, encoding="utf-8")
    logger.info(f"변환 완료: {Path(input_path).name} → {len(text)}자")
    return text


def _clean_pdf_text(text: str) -> str:
    """PDF 추출 후 깨진 문자/제어문자 정제"""
    # 1. PDF 제어문자 제거
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)

    # 2. 깨진 폰트 패턴 제거
    # PDF 헤더/푸터 장식용 폰트: 아랍어/데바나가리/구자라트/말라얄람/버마/키릴 등
    # 연속 2자 이상 나오면 제거 (한글/영문/숫자/기본 특수문자는 보존)
    text = re.sub(
        r"[\u0600-\u06FF\u0700-\u074F\u0900-\u097F\u0A00-\u0A7F"
        r"\u0A80-\u0AFF\u0B00-\u0B7F\u0C00-\u0C7F\u0D00-\u0D7F"
        r"\u0400-\u04FF\u0500-\u052F\u1C00-\u1C4F\u1680-\u169F"
        r"\u07C0-\u07FF\u0750-\u077F\u08A0-\u08FF]{2,}",
        " ", text
    )

    # 3. 연속 공백/빈줄 정리
    text = re.sub(r" {3,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _from_pdf(path: str) -> str:
    reader = PdfReader(path)
    raw = "\n\n".join(
        page.extract_text() or "" for page in reader.pages
    ).strip()
    return _clean_pdf_text(raw)


def _from_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _from_hwpx(path: str) -> str:
    texts = []
    try:
        with zipfile.ZipFile(path) as z:
            section_files = sorted(n for n in z.namelist() if re.search(r"section\d+\.xml", n, re.I))
            targets = section_files or [n for n in z.namelist() if n.endswith(".xml")]
            for fname in targets:
                raw = z.read(fname).decode("utf-8", errors="ignore")
                clean = re.sub(r"<[^>]+>", " ", raw)
                clean = re.sub(r"\s+", " ", clean).strip()
                if clean:
                    texts.append(clean)
    except zipfile.BadZipFile:
        logger.error(f"hwpx 파일 오류: {path}")
    return "\n".join(texts)


def _from_hwp(path: str) -> str:
    soffice = _find_soffice()
    out_dir = str(Path(path).parent)

    try:
        logger.info(f"LibreOffice 변환 시도: {soffice}")
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "txt:Text", path, "--outdir", out_dir],
            capture_output=True,
            text=True,
            timeout=120,
            env=os.environ.copy(),
        )
        txt_path = Path(path).with_suffix(".txt")
        if result.returncode == 0 and txt_path.exists() and txt_path.stat().st_size > 0:
            text = txt_path.read_text(encoding="utf-8", errors="ignore")
            txt_path.unlink()
            logger.info(f"LibreOffice 변환 성공: {len(text)}자")
            return text
        else:
            logger.warning(f"LibreOffice 변환 실패 (returncode={result.returncode}): {result.stderr[:200]}")
    except FileNotFoundError:
        logger.warning(f"LibreOffice 실행파일 없음: {soffice}")
    except subprocess.TimeoutExpired:
        logger.warning("LibreOffice 변환 타임아웃")
    except Exception as e:
        logger.warning(f"LibreOffice 오류: {e}")

    logger.warning("LibreOffice 없음, olefile fallback 시도")
    return _from_hwp_olefile(path)


def _from_hwp_olefile(path: str) -> str:
    try:
        import olefile, zlib
        ole = olefile.OleFileIO(path)
        texts = []
        for stream in ole.listdir():
            if "BodyText" in stream:
                data = ole.openstream(stream).read()
                try:
                    data = zlib.decompress(data, -15)
                except Exception:
                    pass
                decoded = data.decode("utf-16-le", errors="ignore")
                clean = re.sub(r"[^\uAC00-\uD7A3\u0020-\u007E\n]", " ", decoded)
                texts.append(re.sub(r"\s+", " ", clean).strip())
        return "\n".join(texts)
    except Exception as e:
        logger.error(f"HWP olefile 파싱 실패: {e}")
        return ""